from __future__ import annotations

from docx import Document


def extract_docx_text(file_path: str) -> dict:
    document = Document(file_path)

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]

    table_rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_rows.append(cells)

    table_text = "\n".join(" | ".join(row) for row in table_rows)
    joined = "\n".join(paragraphs)
    full_text = "\n\n".join([t for t in [joined, table_text] if t]).strip()

    return {
        "text": full_text,
        "used_ocr": False,
        "tables": table_rows,
    }
