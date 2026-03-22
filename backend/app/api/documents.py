"""
Document Upload + AI Extraction API — Phase 2
POST /documents/upload          → upload → OCR → spaCy → store result
GET  /documents                 → list all uploads
"""
import os
import uuid
import aiofiles
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from uuid import UUID
from typing import List

from app.database import get_db, AsyncSessionLocal
from app.core.dependencies import get_current_user
from app.config import settings
from app.models.user import User
from app.models import Document, Assignment, TaskType, ExtractionStatus
from app.schemas.schemas import DocumentOut, AssignmentOut
from app.ai.extractor import extract_from_text
from app.services.time_estimator import estimate_assignment_time
from app.services import assignment_service

router = APIRouter(prefix="/documents", tags=["Documents"])

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "text/plain": "txt",
    "application/msword": "doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "doc",
}

EXTENSION_TYPES = {
    "pdf": "pdf",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "txt": "txt",
    "doc": "doc",
    "docx": "doc",
}


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF with OCR fallback for scanned/mixed PDFs."""
    try:
        from extractor.pdf_extractor import extract_pdf_text

        result = extract_pdf_text(file_path)
        text = (result or {}).get("text") or ""
        return text
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def _extract_text_from_image(file_path: str) -> str:
    """Extract text from image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
        
        # Configure Tesseract path if specified
        if settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
            
        img = Image.open(file_path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[OCR error: {e}]"


def _extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOCX (and best-effort DOC fallback)."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass

    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        decoded = raw.decode("latin-1", errors="ignore")
        lines = [line.strip() for line in decoded.splitlines()]
        candidates = [line for line in lines if len(line.split()) >= 3]
        return "\n".join(candidates[:400])
    except Exception as e:
        return f"[DOC extraction error: {e}]"


async def _process_document(doc_id: str, user_id: str):
    """
    Background task: extract text → run spaCy → save assignments.
    """
    from uuid import UUID

    async with AsyncSessionLocal() as db:
        try:
            result = await db.execute(select(Document).where(Document.id == UUID(doc_id)))
            doc = result.scalar_one_or_none()
            if not doc:
                return

            doc.extraction_status = ExtractionStatus.processing
            await db.commit()

            # ── Step 1: Extract raw text ──────────────────────────────────
            import anyio
            if doc.file_type == "pdf":
                raw_text = await anyio.to_thread.run_sync(_extract_text_from_pdf, doc.file_path)
            elif doc.file_type == "image":
                raw_text = await anyio.to_thread.run_sync(_extract_text_from_image, doc.file_path)
            elif doc.file_type == "doc":
                raw_text = await anyio.to_thread.run_sync(_extract_text_from_doc, doc.file_path)
            else:  # txt
                async with aiofiles.open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw_text = await f.read()

            doc.raw_text = raw_text[:50000]

            # ── Step 2: NLP extraction ────────────────────────────────────
            extracted_item = extract_from_text(raw_text)

            extracted_task_type = extracted_item.task_type if extracted_item else None
            time_estimate = await estimate_assignment_time(
                text=raw_text,
                task_type=extracted_task_type,
            )

            doc.extracted_data = {
                "title": extracted_item.title if extracted_item else None,
                "subject": extracted_item.subject if extracted_item else None,
                "task_type": extracted_task_type,
                "deadline": str(extracted_item.deadline) if extracted_item and extracted_item.deadline else None,
                "confidence": extracted_item.confidence if extracted_item else 0.0,
                "time_estimate": time_estimate,
            }

            if extracted_item and extracted_item.confidence > 0.2:
                # Map task_type string to enum
                try:
                    task_type = TaskType(extracted_item.task_type)
                except ValueError:
                    task_type = TaskType.assignment

                from app.models.assignment import Assignment
                # Parse deadline (extractors return ISO string) to a date object
                deadline_val = extracted_item.deadline
                if isinstance(deadline_val, str):
                    try:
                        from dateutil import parser as dateutil_parser
                        deadline_val = dateutil_parser.parse(deadline_val).date()
                    except Exception:
                        deadline_val = None

                # Create the official Assignment record
                assignment = Assignment(
                    user_id=UUID(user_id),
                    title=extracted_item.title or "Extracted Assignment",
                    subject=extracted_item.subject,
                    task_type=task_type,
                    deadline=deadline_val,
                    source_document_id=doc.id,
                    ai_metadata={"time_estimate": time_estimate},
                )
                db.add(assignment)

            doc.extraction_status = ExtractionStatus.done
            await db.commit()

        except Exception as e:
            doc.extraction_status = ExtractionStatus.failed
            doc.extraction_error = str(e)
            await db.commit()
            print(f"[extractor] Failed for doc {doc_id}: {e}")


@router.post("/upload", response_model=DocumentOut, status_code=202)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Accepts PDF, image (PNG/JPEG), or TXT file.
    """
    content_type = file.content_type or ""
    file_category = ALLOWED_TYPES.get(content_type)

    if not file_category and file.filename and "." in file.filename:
        ext = file.filename.rsplit(".", 1)[-1].lower()
        file_category = EXTENSION_TYPES.get(ext)

    if not file_category:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type: {content_type}. Allowed: PDF, DOC, DOCX, PNG, JPEG, TXT"
        )

    # Ensure upload directory exists
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    # Save file to disk
    ext = file.filename.rsplit(".", 1)[-1] if "." in file.filename else "bin"
    stored_name = f"{uuid.uuid4()}.{ext}"
    file_path = os.path.join(settings.UPLOAD_DIR, stored_name)
    file_bytes = await file.read()

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(file_bytes)

    # Save document record and commit so the background task can find it
    doc = Document(
        user_id=user.id,
        original_filename=file.filename,
        file_path=file_path,
        file_type=file_category,
        extraction_status=ExtractionStatus.pending,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    # Kick off background extraction via FastAPI BackgroundTasks (reliable)
    background_tasks.add_task(_process_document, str(doc.id), str(user.id))

    return doc


@router.post("/{doc_id}/re-process", status_code=202)
async def reprocess_document(
    doc_id: UUID,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Re-trigger text extraction and assignment creation for a document."""
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    doc.extraction_status = ExtractionStatus.pending
    doc.extraction_error = None
    await db.commit()

    background_tasks.add_task(_process_document, str(doc.id), str(user.id))
    return {"message": "Re-processing started", "doc_id": str(doc_id)}


@router.get("", response_model=List[DocumentOut])
async def list_documents(
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document)
        .where(Document.user_id == user.id)
        .order_by(Document.created_at.desc())
    )
    return result.scalars().all()


@router.get("/{doc_id}", response_model=DocumentOut)
async def get_document(
    doc_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.post("/{doc_id}/save-as-assignment", response_model=AssignmentOut)
async def save_as_assignment_endpoint(
    doc_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """
    Manually trigger conversion of extracted data to an official Assignment record.
    """
    assignment = await assignment_service.create_assignment_from_document(user.id, doc_id, db)
    if not assignment:
        raise HTTPException(status_code=400, detail="Could not create assignment from document. Ensure the document has extracted data.")
    await db.commit()
    await db.refresh(assignment)
    return assignment
